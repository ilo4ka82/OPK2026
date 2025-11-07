"""
Загрузчик и парсер документов для AI-помощника.
Поддерживает: PDF, DOCX, TXT
"""
import os
import logging
from pathlib import Path
from typing import List, Dict
from dataclasses import dataclass

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Класс для хранения чанка документа"""
    text: str
    source: str  # Путь к файлу
    page: int = None  # Номер страницы (для PDF)
    metadata: Dict = None


class DocumentLoader:
    """Загрузчик документов из папки ai_knowledge"""
    
    def __init__(self, knowledge_dir: str = None):
        """
        Args:
            knowledge_dir: Путь к папке с документами.
                          По умолчанию: AI_helper/data/ai_knowledge/
        """
        if knowledge_dir is None:
            # Автоматически определяем путь
            current_dir = Path(__file__).parent
            knowledge_dir = current_dir / "data" / "ai_knowledge"
        
        self.knowledge_dir = Path(knowledge_dir)
        
        if not self.knowledge_dir.exists():
            self.knowledge_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"Создана папка для документов: {self.knowledge_dir}")
    
    def load_all_documents(self, chunk_size: int = 500, chunk_overlap: int = 50) -> List[DocumentChunk]:
        """
        Загружает все документы из папки ai_knowledge
        
        Args:
            chunk_size: Размер чанка в символах
            chunk_overlap: Перекрытие между чанками (для контекста)
        
        Returns:
            Список DocumentChunk
        """
        all_chunks = []
        
        # Рекурсивно ищем все файлы
        for file_path in self.knowledge_dir.rglob("*"):
            if file_path.is_file():
                try:
                    chunks = self._load_file(file_path, chunk_size, chunk_overlap)
                    all_chunks.extend(chunks)
                    logger.info(f"Загружено {len(chunks)} чанков из {file_path.name}")
                except Exception as e:
                    logger.error(f"Ошибка загрузки {file_path}: {e}")
        
        logger.info(f"Всего загружено {len(all_chunks)} чанков из документов")
        return all_chunks
    
    def _load_file(self, file_path: Path, chunk_size: int, chunk_overlap: int) -> List[DocumentChunk]:
        """Загружает один файл в зависимости от расширения"""
        
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return self._load_pdf(file_path, chunk_size, chunk_overlap)
        elif suffix == ".docx":
            return self._load_docx(file_path, chunk_size, chunk_overlap)
        elif suffix == ".txt":
            return self._load_txt(file_path, chunk_size, chunk_overlap)
        else:
            logger.warning(f"Неподдерживаемый формат: {suffix}")
            return []
    
    def _load_pdf(self, file_path: Path, chunk_size: int, chunk_overlap: int) -> List[DocumentChunk]:
        """Загружает PDF файл"""
        chunks = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                text = page.extract_text()
                
                if text.strip():
                    # Разбиваем страницу на чанки
                    page_chunks = self._split_text(text, chunk_size, chunk_overlap)
                    
                    for chunk_text in page_chunks:
                        chunks.append(DocumentChunk(
                            text=chunk_text,
                            source=str(file_path),
                            page=page_num,
                            metadata={
                                "file_name": file_path.name,
                                "file_type": "pdf",
                                "total_pages": len(pdf_reader.pages)
                            }
                        ))
        
        return chunks
    
    def _load_docx(self, file_path: Path, chunk_size: int, chunk_overlap: int) -> List[DocumentChunk]:
        """Загружает DOCX файл"""
        chunks = []
        
        doc = Document(file_path)
        
        # Собираем весь текст
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Разбиваем на чанки
        text_chunks = self._split_text(full_text, chunk_size, chunk_overlap)
        
        for chunk_text in text_chunks:
            chunks.append(DocumentChunk(
                text=chunk_text,
                source=str(file_path),
                metadata={
                    "file_name": file_path.name,
                    "file_type": "docx",
                    "total_paragraphs": len(doc.paragraphs)
                }
            ))
        
        return chunks
    
    def _load_txt(self, file_path: Path, chunk_size: int, chunk_overlap: int) -> List[DocumentChunk]:
        """Загружает TXT файл"""
        chunks = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Разбиваем на чанки
        text_chunks = self._split_text(text, chunk_size, chunk_overlap)
        
        for chunk_text in text_chunks:
            chunks.append(DocumentChunk(
                text=chunk_text,
                source=str(file_path),
                metadata={
                    "file_name": file_path.name,
                    "file_type": "txt"
                }
            ))
        
        return chunks
    
    def _split_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """
        Разбивает текст на чанки с перекрытием
        
        Args:
            text: Исходный текст
            chunk_size: Размер чанка в символах
            chunk_overlap: Перекрытие в символах
        
        Returns:
            Список чанков
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # Пытаемся разорвать на границе предложения
            if end < len(text):
                # Ищем ближайшую точку
                last_period = chunk.rfind(". ")
                if last_period > chunk_size // 2:  # Если точка не слишком далеко
                    end = start + last_period + 1
                    chunk = text[start:end]
            
            if chunk.strip():
                chunks.append(chunk.strip())
            
            # Сдвигаемся с учетом перекрытия
            start = end - chunk_overlap
            
            if start >= len(text):
                break
        
        return chunks


# === ТЕСТИРОВАНИЕ ===
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    loader = DocumentLoader()
    documents = loader.load_all_documents(chunk_size=500, chunk_overlap=50)
    
    print(f"\n📚 Загружено документов: {len(documents)}")
    
    if documents:
        print(f"\n📄 Пример первого чанка:")
        print(f"Источник: {documents[0].source}")
        print(f"Файл: {documents[0].metadata.get('file_name')}")
        if documents[0].page:
            print(f"Страница: {documents[0].page}")
        print(f"Текст: {documents[0].text[:200]}...")
    else:
        print("\n⚠️ Документы не найдены!")
        print(f"Положите PDF/DOCX/TXT файлы в: {loader.knowledge_dir}")